"""Genera el email de lanzamiento del Discovery LP&T como .docx editable en Word."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Márgenes
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_header(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x4A, 0x53, 0x4A)
    p.paragraph_format.space_after = Pt(2)


def add_meta(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.bold = True
    p.add_run(value)
    p.paragraph_format.space_after = Pt(2)


def add_body(text, space_after=10):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)


def add_bullet(text):
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(4)


def add_divider():
    p = doc.add_paragraph()
    r = p.add_run('—' * 40)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ============== EMAIL PRINCIPAL ==============
title = doc.add_paragraph()
tr = title.add_run('Email principal — Lanzamiento Discovery LP&T')
tr.bold = True
tr.font.size = Pt(14)
tr.font.color.rgb = RGBColor(0x1F, 0x3D, 0x2B)
title.paragraph_format.space_after = Pt(14)

add_meta('Para: ', 'martica@..., katherine.monsalve@..., william.nieto@..., jeffrey@..., john.jairo@..., ismelda@..., sergio@..., jhonatan.benavides@..., elisa@..., laura@...')
add_meta('CC: ', 'katherine.rodriguez@...')
add_meta('Asunto: ', 'Necesito tu ayuda — 30 a 60 minutos en los próximos 5 días')

add_divider()

add_body('Equipo,')

add_body('Estamos arrancando un proyecto importante: construir un sistema de asistencia ejecutiva con inteligencia artificial para La Palma & El Tucán. La meta es clara — que el día a día de cada uno sea más fácil, no más complejo.')

add_body('Pero antes de construir, necesitamos entender cómo trabajan ustedes hoy. Por eso preparamos un cuestionario por área. Cada uno tiene el suyo.')

add_body('Acá está todo:')

# Link destacado
link_p = doc.add_paragraph()
link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
link_r = link_p.add_run('👉  https://fsardi19.github.io/lpt-discovery/')
link_r.bold = True
link_r.font.size = Pt(13)
link_r.font.color.rgb = RGBColor(0x1F, 0x3D, 0x2B)
link_p.paragraph_format.space_before = Pt(8)
link_p.paragraph_format.space_after = Pt(14)

add_body('Al entrar, encuentras tu nombre / área. Antes de responder hay un video corto (3 min) que explica el contexto y cómo aprovechar bien el tiempo. Por favor véanlo.')

# QUIÉN RESPONDE QUÉ
add_header('QUIÉN RESPONDE QUÉ')

assignments = [
    ('Martica + equipo contable LP&T', 'Contable'),
    ('Katherine Monsalve', 'Tesorería y Bancos'),
    ('William Nieto', 'Planeación Financiera y Presupuestos'),
    ('Jeffrey', 'Comercial'),
    ('John Jairo', 'Operaciones de Café + Cultivos Asociados'),
    ('Ismelda + Sergio', 'Calidades e Inventarios'),
    ('Jhonatan Benavides', 'Tecnología'),
    ('Elisa + Laura', 'Mercadeo y Comunicación'),
]

for who, area in assignments:
    p = doc.add_paragraph()
    r1 = p.add_run(f'   {who}  →  ')
    r2 = p.add_run(area)
    r2.bold = True
    p.paragraph_format.space_after = Pt(3)

doc.add_paragraph()

# TRES COSAS IMPORTANTES
add_header('TRES COSAS IMPORTANTES')

p = doc.add_paragraph()
r = p.add_run('1. Esto NO es una auditoría. ')
r.bold = True
p.add_run('Es discovery para construir mejor. Si algo está mal hoy, decirlo nos ayuda.')
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
r = p.add_run('2. Sean específicos. ')
r.bold = True
p.add_run('Datos, ejemplos, nombres de archivos — eso vale oro. Las descripciones generales valen poco.')
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
r = p.add_run('3. Hay una sección al final donde piden listar los archivos ')
r.bold = True
p.add_run('que usan a diario (Excel, planillas, etc.). Si manejan muchos, listen los principales 10-15. El resto lo conversamos después.')
p.paragraph_format.space_after = Pt(14)

# PLAZO
plazo = doc.add_paragraph()
r1 = plazo.add_run('PLAZO: ')
r1.bold = True
r1.font.color.rgb = RGBColor(0xA3, 0x7B, 0x3E)
plazo.add_run('5 días hábiles desde hoy. El formulario se guarda solo en el navegador, así que pueden cerrarlo y volver mañana sin perder lo que escribieron.')
plazo.paragraph_format.space_after = Pt(14)

add_body('Cualquier duda, me escriben directamente.')

# Firma
firma = doc.add_paragraph()
firma_r = firma.add_run('Felipe')
firma_r.italic = True
firma.paragraph_format.space_before = Pt(10)


# ============== WHATSAPP RECORDATORIO ==============
add_divider()

title2 = doc.add_paragraph()
tr2 = title2.add_run('WhatsApp / recordatorio (día 3)')
tr2.bold = True
tr2.font.size = Pt(14)
tr2.font.color.rgb = RGBColor(0x1F, 0x3D, 0x2B)
title2.paragraph_format.space_before = Pt(14)
title2.paragraph_format.space_after = Pt(10)

wa = doc.add_paragraph()
wa.add_run('Hola, recordatorio del cuestionario que les mandé el lunes — quedan 2 días hábiles para responder.\n\n')
r = wa.add_run('Link: ')
r.bold = True
wa.add_run('https://fsardi19.github.io/lpt-discovery/\n')
r = wa.add_run('Tiempo: ')
r.bold = True
wa.add_run('30-60 min según área\n\n')
wa.add_run('Si están atrancados con algo, díganme.\nF.')


# ============== TIPS ==============
add_divider()

title3 = doc.add_paragraph()
tr3 = title3.add_run('Tips antes de mandar')
tr3.bold = True
tr3.font.size = Pt(14)
tr3.font.color.rgb = RGBColor(0x1F, 0x3D, 0x2B)
title3.paragraph_format.space_before = Pt(14)
title3.paragraph_format.space_after = Pt(8)

tips = [
    'Verifica que el link se vea completo y clickeable en el preview antes de enviar.',
    'Manda lunes en la mañana (no domingo en la noche, no viernes en la tarde) — la atención es mejor.',
    'WhatsApp paralelo: un mensaje breve al grupo del equipo avisando "acabo de mandar email importante" sube la tasa de apertura ~30%.',
    'Si pasa el día 5 y alguno no responde: llamada directa, no email.',
]

for t in tips:
    p = doc.add_paragraph(f'•  {t}')
    p.paragraph_format.space_after = Pt(4)


# Guardar
out_path = '/Users/felipesardi/Desktop/PERSONAL/AI_AGENTS/openclaw-control/discovery/Email-Lanzamiento-Discovery-LPT.docx'
doc.save(out_path)
print(out_path)
